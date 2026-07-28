"""Dump text from all docx files with unicode-safe output."""
import os, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))

for docx_path in [
    os.path.join(ROOT, "Assets", "Links", "Links.docx"),
    os.path.join(ROOT, "Assets", "Experience_Projects_Volunteering", "Experience_Volunteering_Projects.docx"),
    os.path.join(ROOT, "Assets", "Resume", "ZOHAIB_RAHIM_Master_Resume (2).docx"),
]:
    print("=" * 80)
    print(docx_path)
    print("=" * 80)
    d = Document(docx_path)
    paras = [p.text for p in d.paragraphs if p.text.strip()]
    for i, p in enumerate(paras):
        print(f"{i:03d}: {p}")
    # Also dump tables
    for ti, t in enumerate(d.tables):
        print(f"\n--- TABLE {ti} ---")
        for ri, row in enumerate(t.rows):
            cells = [c.text.strip().replace("\n", " | ") for c in row.cells]
            print(f"  row {ri}: {' | '.join(cells)}")
    print()
